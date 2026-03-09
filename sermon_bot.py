#!/usr/bin/env python3
"""
동감교회 설교 자동 전사 시스템
- Vimeo 채널에서 최신 설교 영상 자동 감지 (찬양 영상 제외)
- 음성 → 텍스트 변환 (faster-whisper, 무료 오픈소스)
- 설교 요약 (OpenAI GPT)
- Word 문서(.docx) 생성
- Google Drive 자동 업로드 (OAuth2, 날짜별 폴더)

영상 제목 패턴:
  설교: "동암교회 OOO 담임목사 - 2026 주일예배"
  찬양: "동암교회 OOO 담임목사 - 2026 주일예배 찬양"
  → "주일예배"가 포함되고 "찬양"이 없는 영상만 선택
"""

import os
import sys
import json
import subprocess
from datetime import datetime

# ============================================================
# 1단계: Vimeo 채널에서 최신 설교 음성 다운로드
# ============================================================
def download_latest_sermon(vimeo_url, output_dir):
    print("=" * 60)
    print("📥 1단계: Vimeo 채널에서 최신 설교 영상 감지 및 다운로드")
    print("=" * 60)

    # 채널의 최근 영상 10개의 제목과 ID를 가져온다
    print("  🔍 채널에서 최근 영상 목록 확인 중...")
    list_cmd = [
        "yt-dlp",
        "--playlist-items", "1:10",
        "--print", "%(title)s|||%(id)s",
        "--no-download",
        vimeo_url
    ]

    result = subprocess.run(list_cmd, capture_output=True, text=True)

    if result.returncode != 0 or not result.stdout.strip():
        print(f"  ❌ 채널 영상 목록을 가져올 수 없습니다!")
        print(f"  stderr: {result.stderr}")
        sys.exit(1)

    # 제목 기반 필터링: "주일예배"가 포함되고 "찬양"이 없는 첫 번째 영상
    lines = [line.strip() for line in result.stdout.strip().split("\n") if line.strip()]

    sermon_title = None
    sermon_id = None

    print(f"  📋 최근 영상 목록:")
    for line in lines:
        if "|||" not in line:
            continue
        title, video_id = line.rsplit("|||", 1)
        title = title.strip()
        video_id = video_id.strip()

        is_sermon = "주일예배" in title and "찬양" not in title
        marker = "✅ 설교" if is_sermon else "⬜ 제외"
        print(f"    {marker}: {title} (ID: {video_id})")

        if is_sermon and sermon_title is None:
            sermon_title = title
            sermon_id = video_id

    if not sermon_title or not sermon_id:
        print("  ❌ 설교 영상을 찾을 수 없습니다! (주일예배 제목 영상 없음)")
        sys.exit(1)

    video_url = f"https://vimeo.com/{sermon_id}"
    print(f"\n  📌 선택된 설교 영상: {sermon_title}")
    print(f"  🔗 영상 URL: {video_url}")

    # 오디오 다운로드
    safe_title = "".join(c for c in sermon_title if c not in r'\/:*?"<>|')[:50]
    today = datetime.now().strftime("%Y%m%d")
    audio_filename = f"설교음성_{today}_{safe_title}"

    download_cmd = [
        "yt-dlp",
        "--extract-audio",
        "--audio-format", "mp3",
        "--audio-quality", "5",
        "--output", os.path.join(output_dir, f"{audio_filename}.%(ext)s"),
        "--no-overwrites",
        video_url
    ]

    print(f"  📥 오디오 다운로드 시작...")
    dl_result = subprocess.run(download_cmd, capture_output=True, text=True)
    print(dl_result.stdout)

    if dl_result.returncode != 0:
        print(f"  ❌ 다운로드 에러: {dl_result.stderr}")
        sys.exit(1)

    audio_path = os.path.join(output_dir, f"{audio_filename}.mp3")
    print(f"  ✅ 다운로드 완료: {audio_path}")

    return audio_path, sermon_title


# ============================================================
# 2단계: 음성 → 텍스트 변환 (Whisper)
# ============================================================
def transcribe_audio(audio_path):
    print("\n" + "=" * 60)
    print("🎙️ 2단계: 음성 → 텍스트 변환 (Whisper)")
    print("=" * 60)

    from faster_whisper import WhisperModel

    print("  모델 로딩 중... (처음 실행 시 다운로드됨)")
    model = WhisperModel("base", device="cpu", compute_type="int8")

    print("  전사 시작... (30분 설교 기준 약 10~20분 소요)")
    segments, info = model.transcribe(audio_path, language="ko")

    full_text = ""
    segment_count = 0
    for segment in segments:
        full_text += segment.text + " "
        segment_count += 1
        if segment_count % 50 == 0:
            print(f"  진행 중... {segment_count}개 세그먼트 처리됨")

    full_text = full_text.strip()
    print(f"  ✅ 전사 완료! 총 {len(full_text)}자, {segment_count}개 세그먼트")

    return full_text


# ============================================================
# 3단계: 설교 요약 (OpenAI GPT)
# ============================================================
def summarize_sermon(transcript, title):
    print("\n" + "=" * 60)
    print("📝 3단계: 설교 요약 (OpenAI GPT)")
    print("=" * 60)

    from openai import OpenAI

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("  ❌ OPENAI_API_KEY 환경변수가 설정되지 않았습니다!")
        sys.exit(1)

    client = OpenAI(api_key=api_key)

    max_chars = 15000
    transcript_trimmed = transcript[:max_chars]
    if len(transcript) > max_chars:
        print(f"  ⚠️ 전사 텍스트가 길어서 앞부분 {max_chars}자만 요약합니다.")

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": """당신은 기독교 설교를 정리하고 요약하는 전문가입니다.
다음 형식으로 설교를 요약해주세요:

[설교 제목]
[본문 말씀] - 성경 구절 (설교 내용에서 파악)
[핵심 메시지] - 3~5줄로 정리
[주요 포인트] - 3~5개 항목
[적용 및 결단] - 2~3줄
[인상 깊은 문장] - 설교에서 인상 깊은 문장 2~3개

한국어로 작성해주세요. 깔끔하고 읽기 쉽게 정리해주세요."""
            },
            {
                "role": "user",
                "content": f"다음 설교를 요약해주세요.\n\n제목: {title}\n\n설교 내용:\n{transcript_trimmed}"
            }
        ],
        max_tokens=2000
    )

    summary = response.choices[0].message.content
    print("  ✅ 요약 완료!")

    return summary


# ============================================================
# 4단계: Word 문서 생성
# ============================================================
def create_word_document(content, filename, title, doc_type="transcript"):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(datetime.now().strftime("%Y년 %m월 %d일 주일설교"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(11)

    doc.add_paragraph("")

    if doc_type == "transcript":
        sentences = content.replace(". ", ".\n").split("\n")
        current_para = ""
        for sent in sentences:
            current_para += sent.strip() + " "
            if len(current_para) > 200:
                p = doc.add_paragraph(current_para.strip())
                p.style.font.size = Pt(11)
                current_para = ""
        if current_para.strip():
            p = doc.add_paragraph(current_para.strip())
            p.style.font.size = Pt(11)
    else:
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("[") and line.endswith("]"):
                doc.add_heading(line[1:-1], level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif len(line) > 1 and line[0].isdigit() and line[1] in ".)" :
                doc.add_paragraph(line[2:].strip(), style="List Number")
            else:
                doc.add_paragraph(line)

    doc.save(filename)
    print(f"  📄 문서 저장 완료: {filename}")


# ============================================================
# 5단계: Google Drive 업로드 (OAuth2, 날짜별 폴더)
# ============================================================
def get_or_create_date_folder(service, parent_folder_id, folder_name):
    """날짜별 하위 폴더를 찾거나 새로 만듭니다."""
    query = (
        f"name='{folder_name}' and "
        f"'{parent_folder_id}' in parents and "
        f"mimeType='application/vnd.google-apps.folder' and "
        f"trashed=false"
    )
    results = service.files().list(q=query, fields="files(id, name)").execute()
    files = results.get("files", [])

    if files:
        folder_id = files[0]["id"]
        print(f"  📁 기존 폴더 사용: {folder_name}")
    else:
        folder_metadata = {
            "name": folder_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [parent_folder_id]
        }
        folder = service.files().create(
            body=folder_metadata,
            fields="id, name"
        ).execute()
        folder_id = folder["id"]
        print(f"  📁 새 폴더 생성: {folder_name}")

    return folder_id


def upload_to_google_drive(file_paths, parent_folder_id):
    print("\n" + "=" * 60)
    print("☁️ 5단계: Google Drive 업로드")
    print("=" * 60)

    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload

    client_id = os.environ.get("GOOGLE_CLIENT_ID")
    client_secret = os.environ.get("GOOGLE_CLIENT_SECRET")
    refresh_token = os.environ.get("GOOGLE_REFRESH_TOKEN")

    if not all([client_id, client_secret, refresh_token]):
        print("  ❌ Google OAuth 환경변수가 설정되지 않았습니다!")
        return False

    try:
        credentials = Credentials(
            token=None,
            refresh_token=refresh_token,
            client_id=client_id,
            client_secret=client_secret,
            token_uri="https://oauth2.googleapis.com/token"
        )
        credentials.refresh(Request())
        service = build("drive", "v3", credentials=credentials)
        print("  ✅ Google Drive 인증 성공!")
    except Exception as e:
        print(f"  ❌ 인증 에러: {e}")
        return False

    # 날짜별 하위 폴더 생성
    today = datetime.now().strftime("%Y%m%d")
    date_folder_name = f"{today}_주일설교"
    target_folder_id = get_or_create_date_folder(service, parent_folder_id, date_folder_name)

    # 파일 업로드
    uploaded_files = []
    for file_path in file_paths:
        if not os.path.exists(file_path):
            print(f"  ⚠️ 파일을 찾을 수 없음: {file_path}")
            continue

        file_name = os.path.basename(file_path)

        mime_types = {
            ".mp3": "audio/mpeg",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".txt": "text/plain",
        }
        ext = os.path.splitext(file_name)[1].lower()
        mime_type = mime_types.get(ext, "application/octet-stream")

        file_metadata = {
            "name": file_name,
            "parents": [target_folder_id]
        }

        media = MediaFileUpload(file_path, mimetype=mime_type, resumable=True)

        try:
            file = service.files().create(
                body=file_metadata,
                media_body=media,
                fields="id, name, webViewLink"
            ).execute()

            print(f"  ✅ 업로드 완료: {file_name}")
            print(f"     링크: {file.get('webViewLink', 'N/A')}")
            uploaded_files.append(file)
        except Exception as e:
            print(f"  ❌ 업로드 실패 ({file_name}): {e}")

    print(f"\n  📊 업로드 결과: {len(uploaded_files)}/{len(file_paths)}개 파일 완료")
    print(f"  📁 저장 위치: {date_folder_name}/")
    return True


# ============================================================
# 메인 실행
# ============================================================
def main():
    print("🔔 동감교회 설교 자동 전사 시스템 시작")
    print(f"   실행 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()

    VIMEO_URL = os.environ.get("VIMEO_URL")
    if not VIMEO_URL:
        print("❌ 에러: VIMEO_URL 환경변수가 설정되지 않았습니다!")
        sys.exit(1)

    DRIVE_FOLDER_ID = os.environ.get("DRIVE_FOLDER_ID", "")

    OUTPUT_DIR = "output"
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # --- 1단계: 설교 영상 감지 & 음성 다운로드 ---
    audio_path, sermon_title = download_latest_sermon(VIMEO_URL, OUTPUT_DIR)

    # --- 2단계: 음성 → 텍스트 ---
    transcript = transcribe_audio(audio_path)

    txt_path = os.path.join(OUTPUT_DIR, "transcript_raw.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(transcript)

    # --- 3단계: 요약 ---
    summary = summarize_sermon(transcript, sermon_title)

    # --- 4단계: Word 문서 생성 ---
    today = datetime.now().strftime("%Y%m%d")
    safe_title = "".join(c for c in sermon_title if c not in r'\/:*?"<>|')[:50]

    transcript_file = os.path.join(OUTPUT_DIR, f"설교전사_{today}_{safe_title}.docx")
    create_word_document(transcript, transcript_file, sermon_title, doc_type="transcript")

    summary_file = os.path.join(OUTPUT_DIR, f"설교요약_{today}_{safe_title}.docx")
    create_word_document(summary, summary_file, sermon_title, doc_type="summary")

    # --- 5단계: Google Drive 업로드 ---
    if DRIVE_FOLDER_ID:
        upload_files = [audio_path, transcript_file, summary_file]
        upload_to_google_drive(upload_files, DRIVE_FOLDER_ID)
    else:
        print("\n⚠️ DRIVE_FOLDER_ID가 설정되지 않아 Google Drive 업로드를 건너뜁니다.")

    # === 완료 ===
    print("\n" + "=" * 60)
    print("🎉 모든 작업이 완료되었습니다!")
    print("-" * 60)
    print(f"  🎵 음성 파일:  {audio_path}")
    print(f"  📄 전사본:     {transcript_file}")
    print(f"  📋 요약본:     {summary_file}")
    print(f"  📝 원문 텍스트: {txt_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
