#!/usr/bin/env python3
"""
동감교회 설교 자동 전사 시스템
- Vimeo에서 최신 설교 영상의 음성 다운로드
- 음성 → 텍스트 변환 (faster-whisper, 무료 오픈소스)
- 설교 요약 (OpenAI GPT)
- Word 문서(.docx) 생성
"""

import os
import sys
import subprocess
from datetime import datetime

# ============================================================
# 1단계: Vimeo에서 최신 설교 음성 다운로드
# ============================================================
def download_latest_sermon(vimeo_url, output_dir):
    """Vimeo에서 최신 영상의 음성만 다운로드합니다."""
    print("=" * 60)
    print("📥 1단계: Vimeo에서 최신 설교 음성 다운로드")
    print("=" * 60)

    # yt-dlp로 최신 영상 1개의 오디오만 다운로드
    cmd = [
        "yt-dlp",
        "--playlist-items", "1",       # 가장 최신 영상 1개만
        "--extract-audio",             # 오디오만 추출
        "--audio-format", "mp3",       # MP3 형식으로 변환
        "--audio-quality", "5",        # 적당한 품질 (파일 크기 절약)
        "--output", os.path.join(output_dir, "sermon_audio.%(ext)s"),
        "--no-overwrites",
        vimeo_url
    ]

    print(f"  다운로드 URL: {vimeo_url}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    print(result.stdout)

    if result.returncode != 0:
        print(f"  ❌ 다운로드 에러: {result.stderr}")
        sys.exit(1)

    audio_path = os.path.join(output_dir, "sermon_audio.mp3")
    print(f"  ✅ 다운로드 완료: {audio_path}")

    # 영상 제목 가져오기 (파일명에 사용)
    title_cmd = [
        "yt-dlp",
        "--playlist-items", "1",
        "--get-title",
        vimeo_url
    ]
    title_result = subprocess.run(title_cmd, capture_output=True, text=True)
    title = title_result.stdout.strip() if title_result.returncode == 0 else "주일설교"
    print(f"  📌 영상 제목: {title}")

    return audio_path, title


# ============================================================
# 2단계: 음성 → 텍스트 변환 (Whisper)
# ============================================================
def transcribe_audio(audio_path):
    """faster-whisper를 사용하여 음성을 텍스트로 변환합니다."""
    print("\n" + "=" * 60)
    print("🎙️ 2단계: 음성 → 텍스트 변환 (Whisper)")
    print("=" * 60)

    from faster_whisper import WhisperModel

    # "base" 모델: GitHub Actions(2코어 CPU)에서 실행 가능한 크기
    # 더 정확한 결과를 원하면 "small"로 변경 (단, 시간이 더 오래 걸림)
    print("  모델 로딩 중... (처음 실행 시 다운로드됨)")
    model = WhisperModel("base", device="cpu", compute_type="int8")

    print("  전사 시작... (30분 설교 기준 약 10~20분 소요)")
    segments, info = model.transcribe(audio_path, language="ko")

    full_text = ""
    segment_count = 0
    for segment in segments:
        full_text += segment.text + " "
        segment_count += 1
        if segment_count % 50 == 0:
            print(f"  진행 중... {segment_count}개 세그먼트 처리됨")

    full_text = full_text.strip()
    print(f"  ✅ 전사 완료! 총 {len(full_text)}자, {segment_count}개 세그먼트")

    return full_text


# ============================================================
# 3단계: 설교 요약 (OpenAI GPT)
# ============================================================
def summarize_sermon(transcript, title):
    """OpenAI GPT를 사용하여 설교를 요약합니다."""
    print("\n" + "=" * 60)
    print("📝 3단계: 설교 요약 (OpenAI GPT)")
    print("=" * 60)

    from openai import OpenAI

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("  ❌ OPENAI_API_KEY 환경변수가 설정되지 않았습니다!")
        sys.exit(1)

    client = OpenAI(api_key=api_key)

    # 전사 텍스트가 너무 길면 잘라서 보냄 (GPT 토큰 제한)
    max_chars = 15000
    transcript_trimmed = transcript[:max_chars]
    if len(transcript) > max_chars:
        print(f"  ⚠️ 전사 텍스트가 길어서 앞부분 {max_chars}자만 요약합니다.")

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": """당신은 기독교 설교를 정리하고 요약하는 전문가입니다.
다음 형식으로 설교를 요약해주세요:

[설교 제목]
[본문 말씀] - 성경 구절 (설교 내용에서 파악)
[핵심 메시지] - 3~5줄로 정리
[주요 포인트] - 3~5개 항목
[적용 및 결단] - 2~3줄
[인상 깊은 문장] - 설교에서 인상 깊은 문장 2~3개

한국어로 작성해주세요. 깔끔하고 읽기 쉽게 정리해주세요."""
            },
            {
                "role": "user",
                "content": f"다음 설교를 요약해주세요.\n\n제목: {title}\n\n설교 내용:\n{transcript_trimmed}"
            }
        ],
        max_tokens=2000
    )

    summary = response.choices[0].message.content
    print("  ✅ 요약 완료!")

    return summary


# ============================================================
# 4단계: Word 문서 생성
# ============================================================
def create_word_document(content, filename, title, doc_type="transcript"):
    """Word 문서(.docx)를 생성합니다."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- 제목 ---
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 날짜 ---
    date_para = doc.add_paragraph(datetime.now().strftime("%Y년 %m월 %d일 주일설교"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(11)

    doc.add_paragraph("")  # 빈 줄

    if doc_type == "transcript":
        # === 전사본: 긴 텍스트를 적당한 문단으로 나누기 ===
        sentences = content.replace(". ", ".\n").split("\n")
        current_para = ""
        for sent in sentences:
            current_para += sent.strip() + " "
            # 약 200자마다 문단 나누기 (읽기 편하게)
            if len(current_para) > 200:
                p = doc.add_paragraph(current_para.strip())
                p.style.font.size = Pt(11)
                current_para = ""
        if current_para.strip():
            p = doc.add_paragraph(current_para.strip())
            p.style.font.size = Pt(11)
    else:
        # === 요약본: 줄바꿈 기준으로 처리 ===
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            # 대괄호로 시작하면 소제목으로 처리
            if line.startswith("[") and line.endswith("]"):
                doc.add_heading(line[1:-1], level=2)
            # 마크다운 헤더(##) 처리
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            # 번호 매기기 또는 글머리표 항목
            elif line.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif len(line) > 1 and line[0].isdigit() and line[1] in ".)" :
                doc.add_paragraph(line[2:].strip(), style="List Number")
            else:
                doc.add_paragraph(line)

    doc.save(filename)
    print(f"  📄 문서 저장 완료: {filename}")


# ============================================================
# 메인 실행
# ============================================================
def main():
    print("🔔 동감교회 설교 자동 전사 시스템 시작")
    print(f"   실행 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print()

    # === 환경변수에서 설정 읽기 ===
    VIMEO_URL = os.environ.get("VIMEO_URL")
    if not VIMEO_URL:
        print("❌ 에러: VIMEO_URL 환경변수가 설정되지 않았습니다!")
        print("   GitHub 저장소 Settings → Secrets에서 VIMEO_URL을 추가해주세요.")
        sys.exit(1)

    # 결과물 저장 폴더
    OUTPUT_DIR = "output"
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # --- 1단계: 음성 다운로드 ---
    audio_path, sermon_title = download_latest_sermon(VIMEO_URL, OUTPUT_DIR)

    # --- 2단계: 음성 → 텍스트 ---
    transcript = transcribe_audio(audio_path)

    # 전사 텍스트를 txt 파일로도 저장 (백업용)
    txt_path = os.path.join(OUTPUT_DIR, "transcript_raw.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(transcript)

    # --- 3단계: 요약 ---
    summary = summarize_sermon(transcript, sermon_title)

    # --- 4단계: Word 문서 생성 ---
    today = datetime.now().strftime("%Y%m%d")

    # 파일명에서 특수문자 제거 (Windows 호환)
    safe_title = "".join(c for c in sermon_title if c not in r'\/:*?"<>|')[:50]

    transcript_file = os.path.join(OUTPUT_DIR, f"설교전사_{today}_{safe_title}.docx")
    create_word_document(transcript, transcript_file, sermon_title, doc_type="transcript")

    summary_file = os.path.join(OUTPUT_DIR, f"설교요약_{today}_{safe_title}.docx")
    create_word_document(summary, summary_file, sermon_title, doc_type="summary")

    # === 완료 ===
    print("\n" + "=" * 60)
    print("🎉 모든 작업이 완료되었습니다!")
    print("-" * 60)
    print(f"  🎵 음성 파일:  {audio_path}")
    print(f"  📄 전사본:     {transcript_file}")
    print(f"  📋 요약본:     {summary_file}")
    print(f"  📝 원문 텍스트: {txt_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
